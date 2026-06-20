from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Helper: set paragraph spacing ────────────────────────────────────────────
def set_spacing(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: set cell border ───────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag  = OxmlElement(f'w:{edge}')
        opts = kwargs.get(edge, kwargs.get('all', {}))
        for k, v in opts.items():
            tag.set(qn(f'w:{k}'), str(v))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: add a run with optional bold/italic/color/size ───────────────────
def add_run(para, text, bold=False, italic=False,
            color=None, size=None, font_name='Times New Roman'):
    run = para.add_run(text)
    run.font.name      = font_name
    run.font.bold      = bold
    run.font.italic    = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run

# ── Helper: heading paragraph ─────────────────────────────────────────────────
def add_heading(text, level=1, color='1F3864', size=14, before=12, after=4):
    para = doc.add_paragraph()
    set_spacing(para, before=before, after=after)
    run = add_run(para, text, bold=True, color=color, size=size)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para

# ── Helper: body paragraph ────────────────────────────────────────────────────
def add_body(text, before=0, after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    para.alignment = align
    set_spacing(para, before=before, after=after)
    add_run(para, text, size=11)
    return para

# ── Helper: bullet ────────────────────────────────────────────────────────────
def add_bullet(text, indent=0):
    para = doc.add_paragraph(style='List Bullet')
    set_spacing(para, before=0, after=3)
    para.paragraph_format.left_indent = Cm(1 + indent)
    add_run(para, text, size=11)
    return para

# ── Helper: code block ────────────────────────────────────────────────────────
def add_code(lines):
    for line in lines:
        para = doc.add_paragraph()
        set_spacing(para, before=0, after=0)
        para.paragraph_format.left_indent = Cm(0.5)
        run = para.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
        # light grey background via shading the paragraph (via XML)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F1F5F9')
        pPr.append(shd)

# ── Helper: page break ────────────────────────────────────────────────────────
def add_page_break():
    para = doc.add_paragraph()
    para.add_run().add_break(WD_BREAK.PAGE)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE / HEADER TABLE  (mimicking the original lab doc header)
# ══════════════════════════════════════════════════════════════════════════════

# ── Main title ────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(title_para, before=0, after=4)
add_run(title_para, 'Lab Session 7 (Open-Ended Lab)',
        bold=True, color='1F3864', size=16)

# ── Subtitle ──────────────────────────────────────────────────────────────────
sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(sub_para, before=0, after=10)
add_run(sub_para, 'Open Ended Lab', bold=True, color='2563EB', size=14)

# ── Taxonomy / Marks table ────────────────────────────────────────────────────
tbl = doc.add_table(rows=2, cols=6)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

headers  = ['Blooms Taxonomy','','', 'GAs','', 'Marks']
sub_hdrs = ['P2','C4','A2','GA5','GA2 / GA6','5 + 10 + 5']

for i, (h, s) in enumerate(zip(headers, sub_hdrs)):
    hdr_cell = tbl.rows[0].cells[i]
    sub_cell = tbl.rows[1].cells[i]
    shade_cell(hdr_cell, '1F3864')
    shade_cell(sub_cell, 'DBEAFE')
    hp = hdr_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(hp, h, bold=True, color='FFFFFF', size=10)
    sp = sub_cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sp, s, bold=True, color='1E3A5F', size=10)

doc.add_paragraph()   # small gap

# ══════════════════════════════════════════════════════════════════════════════
# PROJECT TITLE
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(t, before=6, after=6)
add_run(t, 'Title: Customer Record Management System',
        bold=True, color='1F3864', size=13)

# ══════════════════════════════════════════════════════════════════════════════
# MOTIVATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Motivation:', level=1, size=12, color='1F3864', before=8, after=2)
add_body(
    'Discover problem-solving skills, creativity, and a deeper understanding of web '
    'technologies, preparing you for real-world challenges in web development.\n\n'
    'In today\'s digital age, web applications are a cornerstone of communication, '
    'collaboration, and information sharing. From social media to e-commerce platforms, '
    'real-time interactivity and dynamic data handling are crucial for user engagement '
    'and satisfaction. This lab offers the opportunity to bridge the gap between theory '
    'and practice by building modern web applications using HTML5, CSS3, JavaScript & '
    'ASP.NET framework that emphasise dynamic content updates, responsive design, and '
    'efficient data management.'
)

# ══════════════════════════════════════════════════════════════════════════════
# CONCEPT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Concept:', size=12, color='1F3864', before=6, after=2)
add_body(
    'A web-based application that tracks daily spending, categorises expenditures, '
    'and displays summaries to assist users in efficiently managing their money.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Problem Statement:', size=12, color='1F3864', before=6, after=2)
add_body(
    'Due to an insufficiency of easily accessible resources for tracking and evaluating '
    'spending, many people find it difficult to manage their money. The goal of this '
    'project is to create a straightforward, perceptive web application for tracking '
    'and classifying spending in order to provide information about spending patterns.'
)

# ══════════════════════════════════════════════════════════════════════════════
# DESIGN / WAYS & MEANS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Design / Ways & Means:', size=12, color='1F3864', before=6, after=2)

add_heading('Introduction and Requirements:', size=11, color='2563EB', before=4, after=2)
for b in [
    'Develop an intuitive web application that enables users to track their expenditures '
    'in detail, including the date, category, and amount.',
    'See an overview of all expenses, broken down by category.',
    'Edit or remove entries.',
]:
    add_bullet(b)

add_heading('Web Application Design:', size=11, color='2563EB', before=6, after=2)
add_heading('Frontend:', size=11, color='374151', before=2, after=1)
add_bullet('HTML and CSS for a clean, responsive layout.')
add_bullet('JavaScript for interactive and dynamic functionality.')
add_heading('Backend:', size=11, color='374151', before=2, after=1)
add_bullet('ASP.NET Core MVC for data handling and server-side logic.')

add_heading('Basic Implementation:', size=11, color='2563EB', before=6, after=2)
for b in [
    'Make a table to show logged expenses and an expense input form.',
    'Include basic features for calculating totals, editing, and deleting data.',
    'Present summary data, such as total expenditures and expenditures by category.',
]:
    add_bullet(b)

add_heading('Performance Testing and Analysis:', size=11, color='2563EB', before=6, after=2)
add_body('Test for:')
add_bullet('Accurate expense calculations, management and updates.')
add_bullet('Responsive Application.')

add_heading('Extensions and Creativity:', size=11, color='2563EB', before=6, after=2)
for b in [
    'Implement user log-in to save secure expense data.',
    'Introduce a budget-tracking feature with alerts for overspending and max limit.',
    'Allow data export for external analysis.',
]:
    add_bullet(b)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE BREAK → DELIVERABLES
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()

deliv_para = doc.add_paragraph()
deliv_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(deliv_para, before=0, after=8)
add_run(deliv_para, 'Deliverables', bold=True, color='1F3864', size=14)

# ══════════════════════════════════════════════════════════════════════════════
# 1. BACKGROUND / THEORY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1.  Background / Theory', size=13, color='1F3864', before=8, after=4)

add_body(
    'Effective money management is a challenge for many individuals due to the lack of '
    'accessible and user-friendly tools for tracking and analysing daily expenditures. '
    'Without a clear understanding of spending habits, individuals often struggle to '
    'allocate their finances effectively, leading to potential overspending and financial '
    'instability.'
)
add_body(
    'The development of a web-based application for tracking daily spending addresses '
    'this gap by offering a simple and intuitive platform where users can log their '
    'expenses, categorise them, and view summaries. This categorisation and analysis '
    'provide valuable insights into spending patterns, helping users make informed '
    'financial decisions and manage their money more efficiently.'
)
add_body(
    'By utilising technology to simplify the process of financial tracking, the '
    'application empowers users to gain control over their spending, fostering better '
    'financial discipline and awareness.'
)

add_heading('1.1  Technologies Used', size=12, color='2563EB', before=6, after=3)

# Technologies table
tech_tbl = doc.add_table(rows=7, cols=3)
tech_tbl.style = 'Table Grid'
tech_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

tech_headers = ['Technology', 'Role', 'Details']
for i, h in enumerate(tech_headers):
    c = tech_tbl.rows[0].cells[i]
    shade_cell(c, '1F3864')
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, color='FFFFFF', size=10)

tech_data = [
    ('HTML5',          'Markup',        'Semantic structure for all pages and forms'),
    ('CSS3',           'Styling',       'Responsive layout using Flexbox, Grid and CSS Variables'),
    ('JavaScript',     'Interactivity', 'Sidebar toggle, delete confirmation, alert dismissal'),
    ('ASP.NET Core 10','Backend',       'MVC framework, routing, Razor engine, model binding'),
    ('Razor Views',    'Templating',    'Server-side HTML generation with C# expressions'),
    ('In-Memory Store','Data Layer',    'Static C# list — no database required for this lab'),
]
for row_i, (tech, role, detail) in enumerate(tech_data, start=1):
    row = tech_tbl.rows[row_i]
    shade_cell(row.cells[0], 'DBEAFE' if row_i % 2 == 0 else 'EFF6FF')
    for col_i, text in enumerate([tech, role, detail]):
        p = row.cells[col_i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_i < 2 else WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, bold=(col_i == 0), size=10)

doc.add_paragraph()

add_heading('1.2  MVC Architecture', size=12, color='2563EB', before=4, after=2)
add_body(
    'The application follows the Model-View-Controller (MVC) architectural pattern, '
    'which separates concerns into three distinct layers:'
)
add_bullet('Model – C# classes (Expense, ExpenseStore) that define data structure and validation rules.')
add_bullet('View  – Razor .cshtml files that render HTML pages dynamically from model data.')
add_bullet('Controller – C# classes (HomeController, ExpenseController) that handle HTTP requests and return views.')

# ══════════════════════════════════════════════════════════════════════════════
# PAGE BREAK → CODE
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. CODE IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2.  Fully Functional & Well-Documented Code', size=13, color='1F3864', before=4, after=4)

add_heading('2.1  Project Structure', size=12, color='2563EB', before=4, after=2)
add_code([
    'ExpenseTracker/',
    '│',
    '├── Controllers/',
    '│   ├── HomeController.cs        // Dashboard statistics',
    '│   └── ExpenseController.cs     // CRUD – Create / Read / Update / Delete',
    '│',
    '├── Models/',
    '│   ├── Expense.cs               // Data model + Data Annotations',
    '│   ├── ExpenseStore.cs          // In-memory data store (static list)',
    '│   └── CategoryTotal.cs         // DTO for bar-chart data',
    '│',
    '├── Views/',
    '│   ├── _ViewImports.cshtml      // Tag helpers + namespaces',
    '│   ├── _ViewStart.cshtml        // Default layout declaration',
    '│   ├── Shared/_Layout.cshtml    // Sidebar + Topbar shell',
    '│   ├── Home/Index.cshtml        // Dashboard page',
    '│   └── Expense/',
    '│       ├── Index.cshtml         // Expense list + filters',
    '│       ├── Create.cshtml        // Add expense form',
    '│       └── Edit.cshtml          // Edit expense form',
    '│',
    '├── wwwroot/',
    '│   ├── css/site.css             // Full responsive stylesheet',
    '│   └── js/site.js               // Client-side interactivity',
    '│',
    '├── Program.cs                   // App entry point & middleware',
    '└── ExpenseTracker.csproj        // .NET 10 project configuration',
])

doc.add_paragraph()
add_heading('2.2  Expense Data Model  (Models/Expense.cs)', size=12, color='2563EB', before=6, after=2)
add_body('The Expense class defines each record\'s structure. Data Annotations enforce validation automatically by ASP.NET Core model binding:')
add_code([
    'using System.ComponentModel.DataAnnotations;',
    '',
    'namespace ExpenseTracker.Models',
    '{',
    '    public class Expense',
    '    {',
    '        public int Id { get; set; }',
    '',
    '        [Required(ErrorMessage = "Description is required")]',
    '        [StringLength(100)]',
    '        public string Description { get; set; } = string.Empty;',
    '',
    '        [Required(ErrorMessage = "Amount is required")]',
    '        [Range(0.01, 999999.99, ErrorMessage = "Amount must be > 0")]',
    '        [DataType(DataType.Currency)]',
    '        public decimal Amount { get; set; }',
    '',
    '        [Required(ErrorMessage = "Category is required")]',
    '        public string Category { get; set; } = string.Empty;',
    '',
    '        [Required][DataType(DataType.Date)]',
    '        public DateTime Date { get; set; } = DateTime.Today;',
    '',
    '        public string? Notes { get; set; }',
    '    }',
    '}',
])

add_heading('2.3  In-Memory Data Store  (Models/ExpenseStore.cs)', size=12, color='2563EB', before=6, after=2)
add_body('A static list acts as the data layer — no database required for this lab:')
add_code([
    'public static class ExpenseStore',
    '{',
    '    private static int _nextId = 1;',
    '',
    '    public static List<Expense> Expenses { get; } = new()',
    '    {',
    '        new Expense { Id=_nextId++, Description="Grocery shopping",',
    '                      Amount=85.50m, Category="Food",',
    '                      Date=DateTime.Today.AddDays(-1) },',
    '        new Expense { Id=_nextId++, Description="Bus pass",',
    '                      Amount=30.00m, Category="Transport",',
    '                      Date=DateTime.Today.AddDays(-2) },',
    '        // ... additional seed records',
    '    };',
    '',
    '    public static readonly List<string> Categories = new()',
    '    {',
    '        "Food","Transport","Shopping","Entertainment",',
    '        "Bills","Health","Education","Other"',
    '    };',
    '',
    '    public static int GetNextId() => _nextId++;',
    '}',
])

add_page_break()

add_heading('2.4  Expense Controller  (Controllers/ExpenseController.cs)', size=12, color='2563EB', before=4, after=2)
add_body('Handles all CRUD operations. The Index action supports filtering by category, keyword, and date range:')
add_code([
    'public class ExpenseController : Controller',
    '{',
    '    // GET: /Expense  (with optional query filters)',
    '    public IActionResult Index(string? category, string? search,',
    '                               string? dateFrom, string? dateTo)',
    '    {',
    '        var expenses = ExpenseStore.Expenses.AsQueryable();',
    '',
    '        if (!string.IsNullOrEmpty(category) && category != "All")',
    '            expenses = expenses.Where(e => e.Category == category);',
    '',
    '        if (!string.IsNullOrEmpty(search))',
    '            expenses = expenses.Where(e => e.Description',
    '                .Contains(search, StringComparison.OrdinalIgnoreCase));',
    '',
    '        if (DateTime.TryParse(dateFrom, out var from))',
    '            expenses = expenses.Where(e => e.Date >= from);',
    '',
    '        ViewBag.Total = expenses.Sum(e => e.Amount);',
    '        return View(expenses.OrderByDescending(e => e.Date).ToList());',
    '    }',
    '',
    '    // POST: /Expense/Create',
    '    [HttpPost][ValidateAntiForgeryToken]',
    '    public IActionResult Create(Expense expense)',
    '    {',
    '        if (ModelState.IsValid)',
    '        {',
    '            expense.Id = ExpenseStore.GetNextId();',
    '            ExpenseStore.Expenses.Add(expense);',
    '            TempData["Success"] = "Expense added successfully!";',
    '            return RedirectToAction(nameof(Index));',
    '        }',
    '        ViewBag.Categories = ExpenseStore.Categories;',
    '        return View(expense);',
    '    }',
    '',
    '    // POST: /Expense/Edit/{id}',
    '    [HttpPost][ValidateAntiForgeryToken]',
    '    public IActionResult Edit(int id, Expense updated)',
    '    {',
    '        if (ModelState.IsValid)',
    '        {',
    '            var e = ExpenseStore.Expenses.First(x => x.Id == id);',
    '            e.Description = updated.Description;',
    '            e.Amount      = updated.Amount;',
    '            e.Category    = updated.Category;',
    '            e.Date        = updated.Date;',
    '            e.Notes       = updated.Notes;',
    '            TempData["Success"] = "Expense updated!";',
    '            return RedirectToAction(nameof(Index));',
    '        }',
    '        return View(updated);',
    '    }',
    '',
    '    // POST: /Expense/Delete/{id}',
    '    [HttpPost][ValidateAntiForgeryToken]',
    '    public IActionResult Delete(int id)',
    '    {',
    '        var e = ExpenseStore.Expenses.FirstOrDefault(x => x.Id == id);',
    '        if (e != null) ExpenseStore.Expenses.Remove(e);',
    '        TempData["Success"] = "Expense deleted!";',
    '        return RedirectToAction(nameof(Index));',
    '    }',
    '}',
])

add_heading('2.5  CSS Design System  (wwwroot/css/site.css – excerpt)', size=12, color='2563EB', before=6, after=2)
add_body('CSS Custom Properties (variables) drive the entire colour scheme, making global restyling easy:')
add_code([
    ':root {',
    '    --primary:       #2563EB;',
    '    --primary-dark:  #1D4ED8;',
    '    --success:       #10B981;',
    '    --danger:        #EF4444;',
    '    --sidebar-width: 240px;',
    '    --bg:            #F1F5F9;',
    '    --radius:        10px;',
    '}',
    '',
    '/* Fixed sidebar */',
    '.sidebar {',
    '    width: var(--sidebar-width);',
    '    background: #0F172A;',
    '    position: fixed;',
    '    top: 0; left: 0;',
    '    min-height: 100vh;',
    '}',
    '',
    '/* 4-column responsive stat cards */',
    '.cards-grid {',
    '    display: grid;',
    '    grid-template-columns: repeat(4, 1fr);',
    '    gap: 16px;',
    '}',
    '',
    '/* Mobile – collapse sidebar */',
    '@media (max-width: 900px) {',
    '    .sidebar { transform: translateX(-100%); }',
    '    .sidebar.open { transform: translateX(0); }',
    '    .main-wrapper { margin-left: 0; }',
    '}',
])

add_heading('2.6  JavaScript  (wwwroot/js/site.js)', size=12, color='2563EB', before=6, after=2)
add_code([
    '// Toggle sidebar visibility on mobile',
    'function toggleSidebar() {',
    '    document.getElementById("sidebar").classList.toggle("open");',
    '}',
    '',
    '// Confirm before deleting a record',
    'function confirmDelete() {',
    '    return confirm("Are you sure you want to delete this expense?");',
    '}',
    '',
    '// Auto-dismiss success alerts after 4 seconds',
    'document.addEventListener("DOMContentLoaded", function () {',
    '    const alert = document.querySelector(".alert");',
    '    if (alert) {',
    '        setTimeout(() => {',
    '            alert.style.opacity    = "0";',
    '            alert.style.transition = "opacity 0.4s";',
    '            setTimeout(() => alert.remove(), 400);',
    '        }, 4000);',
    '    }',
    '});',
])

# ══════════════════════════════════════════════════════════════════════════════
# PAGE BREAK → SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()

add_heading('3.  Output Screenshots', size=13, color='1F3864', before=4, after=4)

IMAGES_DIR = os.path.join(os.path.dirname(__file__), 'images')

def add_screenshot(filename, caption_text, description):
    add_heading(caption_text, size=12, color='2563EB', before=6, after=2)
    add_body(description)
    img_path = os.path.join(IMAGES_DIR, filename)
    if os.path.exists(img_path):
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(img_path, width=Inches(5.8))
    else:
        # Placeholder box
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(ph, before=4, after=4)
        add_run(ph,
                f'[ Screenshot: {filename} — place image in Report/images/ folder ]',
                italic=True, color='94A3B8', size=10)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(cap, before=3, after=8)
    add_run(cap, caption_text, italic=True, color='64748B', size=10)

add_screenshot(
    'dashboard.png',
    'Figure 1 – Dashboard (Home Page)',
    'The home page displays four summary cards showing total expenses, this month\'s '
    'total, the top spending category, and total transaction count. Below, a horizontal '
    'bar chart breaks down spending by category, and the five most recent expenses are '
    'listed on the right panel.'
)

add_screenshot(
    'add-expense.png',
    'Figure 2 – Add Expense Form',
    'The Add Expense page presents a clean form with fields for Date, Category '
    '(dropdown with 8 options), Description, Amount (with $ prefix), and optional '
    'Notes. All fields are validated server-side; errors are displayed inline beneath '
    'each field.'
)

add_screenshot(
    'all-expenses.png',
    'Figure 3 – All Expenses Table',
    'The Expenses page lists all records in a sortable table. A filter bar at the top '
    'allows filtering by category, keyword search, and date range. Each row includes '
    'a colour-coded category badge and Edit / Delete action buttons. A grand total is '
    'shown in the table footer.'
)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE BREAK → REFLECTION
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()

add_heading('4.  Reflection Report', size=13, color='1F3864', before=4, after=4)

# 4.1 Challenges
add_heading('4.1  Challenges Faced During the Project', size=12, color='2563EB', before=4, after=3)

challenges = [
    ('Razor Expression Syntax',
     'Calling .ToString("N2") outside a cast expression such as '
     '$@((decimal)ViewBag.Total).ToString("N2") caused the format string to render as '
     'literal text on screen. The fix was to wrap the entire expression inside one '
     'Razor block: @("$" + value.ToString("N2")).'),
    ('_ViewImports.cshtml File Location',
     'Placing _ViewImports.cshtml inside Views/Shared/ instead of Views/ caused Tag '
     'Helpers (asp-for, asp-controller, asp-append-version) and using-namespace '
     'directives to fail for all controller views, breaking CSS links and form binding.'),
    ('HTTPS Redirect on HTTP-Only Server',
     'The UseHttpsRedirection() middleware redirected incoming HTTP traffic to an HTTPS '
     'port that was never configured, resulting in an HTTP 403 Access Denied error in '
     'the browser. Removing the middleware resolved the issue for local development.'),
    ('Dynamic ViewBag vs. Typed Classes',
     'Passing anonymous types through ViewBag and accessing them as dynamic in Razor '
     'views caused runtime binding failures. Introducing a concrete CategoryTotal DTO '
     'class provided strongly-typed access and resolved the error.'),
    ('Razor Option "selected" Attribute (RZ1031)',
     'The compiler error RZ1031 — "tag helper must not have C# in the element\'s '
     'attribute declaration area" — occurred when using an inline ternary operator for '
     'the selected attribute. Replacing it with an explicit @if / else block fixed the '
     'issue.'),
    ('Target Framework Mismatch',
     'The project was initially configured to target net8.0 while only the .NET 10 SDK '
     'was installed. This produced a slow workload-resolution warning and a potential '
     'compatibility issue. Updating the .csproj TargetFramework to net10.0 resolved it.'),
]

for i, (title, body) in enumerate(challenges, start=1):
    p = doc.add_paragraph()
    set_spacing(p, before=4, after=2)
    add_run(p, f'{i}.  {title}', bold=True, color='1F3864', size=11)
    add_body(body, before=0, after=6)

# 4.2 Potential Improvements
add_heading('4.2  Potential Improvements (Given More Time)', size=12, color='2563EB', before=6, after=3)
improvements = [
    'Database Integration — Replace the in-memory store with Entity Framework Core '
    '+ SQLite/SQL Server for true data persistence across server restarts.',
    'User Authentication — Add ASP.NET Core Identity so each registered user has '
    'a private, secure set of expense records.',
    'Budget Alerts — Allow users to set monthly limits per category and receive '
    'visual warnings when approaching or exceeding the limit.',
    'Interactive Charts — Integrate Chart.js to render animated pie, doughnut, and '
    'line-trend charts, providing richer visual analytics.',
    'Data Export — Add CSV and PDF export so users can analyse data in Excel or '
    'share it with an accountant.',
    'Advanced Analytics — Display week-over-week and month-over-month spending '
    'trends and predict future expenditure using simple moving averages.',
]
for imp in improvements:
    add_bullet(imp)

# 4.3 Discussion on Results
add_heading('4.3  Discussion on Results', size=12, color='2563EB', before=6, after=3)
add_body(
    'The web-based expense tracker successfully met all primary objectives defined in '
    'the problem statement. The application provides users with a clean, responsive '
    'interface to log daily expenses, assign them to meaningful categories, and view '
    'both detailed records and high-level summaries on the dashboard.'
)
add_body(
    'The category bar chart offers an immediate visual understanding of where money '
    'is being spent, while the filter and search capabilities on the expenses table '
    'allow users to quickly isolate specific transactions by date range or keyword. '
    'Edit and delete functionality ensure records remain accurate over time.'
)
add_body(
    'From a technical standpoint, the project reinforced key ASP.NET Core MVC concepts '
    'including routing conventions, model binding, data annotations for validation, '
    'anti-forgery tokens for security, and the Razor view engine. The CSS design system '
    'built with custom properties and CSS Grid demonstrates how a consistent, scalable '
    'stylesheet can be authored without any third-party UI framework.'
)
add_body(
    'Overall, the application fulfils the brief and provides a solid foundation upon '
    'which real-world features — persistent storage, user accounts, and rich analytics '
    '— could be built in a production context.'
)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_page_break()
add_heading('References', size=13, color='1F3864', before=4, after=6)

refs = [
    'Reid, M. (2020). Developing Web Applications with ASP.NET Core and Entity Framework '
    'Core. Packt Publishing. A comprehensive guide on building web-based applications '
    'using the ASP.NET Core framework.',
    'Bergheim, B. D., Garrett, D. M., & Maki, D. M. (2001). Education and Saving: The '
    'Long-Term Effects of High School Financial Curriculum Mandates. Journal of Public '
    'Economics, 80(3), 435–465.',
    'Microsoft. (2024). ASP.NET Core MVC Documentation. '
    'https://learn.microsoft.com/en-us/aspnet/core/mvc/overview',
    'MDN Web Docs. (2024). CSS Custom Properties (Variables). '
    'https://developer.mozilla.org/en-US/docs/Web/CSS/Using_CSS_custom_properties',
]
for i, ref in enumerate(refs, start=1):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=8)
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    add_run(p, f'[{i}]  ', bold=True, color='2563EB', size=11)
    add_run(p, ref, size=11)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), 'Lab_Report_Expense_Tracker.docx')
doc.save(out)
print(f'Saved → {out}')
